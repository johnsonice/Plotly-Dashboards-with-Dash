
# coding: utf-8

# ### 1. Read in word document, identify "hot button issue" by checking against pre-defined keyword list 
# ### 2. Return dictionary {"hot button issue": set of keywords found}
# 
# ### Note: Keyword needs to include mutations such as " intervene -> intervened, intervines"

# In[1]:


import sys
import os
import pickle
import re
import pandas as pd
import numpy as np
from docx import Document
import time
from collections import Counter
try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile


class Hotbutton_finder(object):
    """
    an hanlp analyzer object for dependency parsing an other related operations 
    """
    def __init__(self,hot_button_file=None,hot_button_dict_path=None):
        self.hot_button_file = hot_button_file
        self.hot_button_dict_path = hot_button_dict_path
        self.hot_button_dict = self.get_topic_keywords_pairs()
        print('Load hotbutton dictionary successfully...')

    ## define a function to loacte keywords
    @staticmethod
    def construct_rex(keywords):
        r_keywords = [r'\b' + re.escape(k) + r'(s|es|d|ed)?\b'for k in keywords]    # tronsform keyWords list to a patten list, find both s and es 
        rex = re.compile('|'.join(r_keywords),flags=re.I)                        # use or to join all of them, ignore casing
        #match = [(m.start(),m.group()) for m in rex.finditer(content)]          # get the position and the word
        return rex
    

    def find_exact_keywords(self,content,keywords,rex=None):
        if rex is None: 
            rex = self.construct_rex(keywords)
        content = content.replace('\n', '').replace('\r', '')#.replace('.',' .')
        match = Counter([m.group() for m in rex.finditer(content)])             # get all instances of matched words 
                                                                                 # and turned them into a counter object, to see frequencies
        return match
    
    @staticmethod
    def consolidate_counts(key_counts,thresh):
        values = sum(key_counts.values())
        if values > thresh:
            return True
        else:
            return False
        
    def transform_dict_to_re_pair(self,temp_hotbutton_dict):
        '''Create dictionary mapping name and regular expression'''
        for k, v in temp_hotbutton_dict.items():
            # regular expression, deal with multiplles
            ## 1. adding 's','d'at end
            ## 2. separated from other characters
            ## 3. take care of duplicated space
            vs = v.split(",")
            ks = [k.strip() for k in vs if len(k)>1]
            temp_hotbutton_dict[k] = self.construct_rex(ks)
    
        return temp_hotbutton_dict
    
    def get_topic_keywords_pairs(self):
        '''read hot button issues and transform to dic of re expression'''
        if self.hot_button_dict_path and os.path.isfile(self.hot_button_dict_path):
            re_hot_button_dict = pickle.load(open(self.hot_button_dict_path, 'rb'))
            print('load from previous saved dict')
            return re_hot_button_dict
        
        ## read raw file
        if os.path.isfile(self.hot_button_file):
            hot_button_df = pd.read_excel(self.hot_button_file)
            hot_button_df.fillna('', inplace= True)
            ## concanate all search words 
            hot_button_df['keyword list'] = hot_button_df['related words selected'].str.cat(hot_button_df['augmented words from topic modelling'], 
                         sep = ', ').str.cat(hot_button_df['augmented words from word2vec'], 
                                   sep = ', ').str.cat(hot_button_df['search term for word2vec'], 
                                             sep = ', ')
            hot_button_df['keyword list'] = hot_button_df['keyword list'].str.lower().str.replace(r'/|-|_',' ')   
            hot_button_dict = pd.Series(hot_button_df['keyword list'].values,index=hot_button_df['Hot button issues']).to_dict()
            
            re_hot_button_dict = self.transform_dict_to_re_pair(hot_button_dict)
            pickle.dump(re_hot_button_dict, open(self.hot_button_dict_path,'wb'))
            print('Generate new hot_button_dict in {}'.format(self.hot_button_dict_path))
        else:
            raise Exception('file does not exist: {}'.format(self.hot_button_file))
            
        return re_hot_button_dict
    
    
    def check_all_topics(self,document,thresh=0):
        res_list = []
        for k, v in self.hot_button_dict.items():
            res = self.find_exact_keywords(document,keywords=None,rex=v)
            #print(res)
            check = self.consolidate_counts(res,thresh)
            if check:
                res_list.append(k)
        return res_list
    
    @staticmethod
    def get_footnotes_text(f_path):
        """
        Take the path of a docx file as argument, return the text in unicode.
        """
        try:
            document = zipfile.ZipFile(f_path)
            xml_content = document.read('word/footnotes.xml')
            document.close()
            tree = XML(xml_content)
            reslist = list(tree.iter())
            results = ' '.join([e.text for e in reslist if e.text is not None])
        except:
            print('no footnotes found')
            return ""
        
        return results
    
    
    def read_doc(self,f_path,word_length_filter=0):
        '''load an clean document'''
        doc = Document(f_path)
        text_list = [p.text for p in doc.paragraphs]#[3:]
        text_list = [p.replace('\xa0',' ') for p in text_list] # some clean up 
        text_list = [p for p in text_list if len(p.split()) > word_length_filter]
        ## some replacement of "-/_" to space
        document = list(map(lambda x: re.sub(r'—|-|_',' ',x), text_list))
        document_text = " ".join(document).lower()
        
        ## get table text
        try:
            table_text = []
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text is not None and len(paragraph.text)>0:
                                table_text.append(paragraph.text)
            table_text = " ".join(table_text)
        except:
            print('no table text found')
            table_text=' '
            
        ## get footnotes 
        footnotes_text = self.get_footnotes_text(f_path)
        
        
        final_text = '{} {} {}'.format(document_text, table_text, footnotes_text)
        #final_text = re.sub(r'—|-|_',' ',final_text)
        
        return final_text


class Custom_finder(Hotbutton_finder):
    def __init__(self,custom_file=None,custom_dict_path=None):
        self.custom_file = custom_file
        self.custom_dict_path = custom_dict_path
        Hotbutton_finder.__init__(self,hot_button_file=custom_file,hot_button_dict_path=custom_dict_path)
        self.hot_button_dict = self.get_topic_keywords_pairs()
        print('Load customized dictionary successfully...')
        
    def get_topic_keywords_pairs(self):
        print('this is a different function')
        '''read hot button issues and transform to dic of re expression'''
        if self.custom_file and os.path.isfile(self.custom_dict_path):
            re_custom_dict = pickle.load(open(self.custom_dict_path, 'rb'))
            print('load from previous saved dict')
            return re_custom_dict
        
        ## read raw file
        if os.path.isfile(self.custom_file):
            custom_df = pd.read_excel(self.custom_file)
            custom_df.fillna('', inplace= True)
            ## concanate all search words 
            custom_df['key_words'] = custom_df['key_words'].str.lower().str.replace(r'/|-|_',' ')   
            custom_dict = pd.Series(custom_df['key_words'].values,index=custom_df['name']).to_dict()
            
            re_custom_dict = self.transform_dict_to_re_pair(custom_dict)
            pickle.dump(re_custom_dict, open(self.custom_dict_path,'wb'))
            print('Generate new custom_dict in {}'.format(self.custom_dict_path))
        else:
            raise Exception('file does not exist: {}'.format(self.custom_file))
# In[5]:

if __name__ == '__main__':
    
    # Save and load for reuse in production
    start_time = time.time()
    
    hot_button_file = os.path.join('./model_weights/keywords_search/hot_button_issues.xlsx')
    #text_file_path = "./test/Brazil_2013.DOCX"
    text_file_path = "./test/Canada-2018.docx"
    save_file = './model_weights/keywords_search/hot_button_dict.pickle'
    print(os.getcwd())
    
    ## initiate hotbutton object 
    hotbutton_finder = Hotbutton_finder(hot_button_file,save_file)
    #hotbutton_finder = Custom_finder(hot_button_file,save_file)
    
    document = hotbutton_finder.read_doc(text_file_path)
    print(hotbutton_finder.check_all_topics(document))

    ## run one test 
    test = 'asd macroprudential asdf savegards  resources reassure discipline government corruption'
    found_topics = hotbutton_finder.check_all_topics(test)
    print(found_topics)
    
    ## test custom dict
    
    custom_file = './model_weights/keywords_search/minimum_requirement.xlsx'
    save_file = './model_weights/keywords_search/minimum_requirement.pickle'
    #text_file_path = "./test/Brazil_2013.DOCX"
    CF = Custom_finder(custom_file,save_file)
    document = CF.read_doc(text_file_path)
    print(CF.check_all_topics(document))
    


    